"""Generate Jakevolume_Architecture.docx — two-page architecture reference."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: landscape 11×8.5 ─────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin   = Inches(0.6)
    section.right_margin  = Inches(0.6)
    section.top_margin    = Inches(0.5)
    section.bottom_margin = Inches(0.5)

NAVY  = RGBColor(0x1A, 0x37, 0x5E)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
GOLD  = RGBColor(0xC8, 0x8B, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LTBLUE = RGBColor(0xD6, 0xE4, 0xF0)
LTGOLD = RGBColor(0xFD, 0xF3, 0xCC)
LTGREEN = RGBColor(0xD6, 0xF0, 0xDC)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def hdr_para(text, size=22, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    return p


def sub_para(text, size=9, italic=False, color=RGBColor(0x33, 0x33, 0x33)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    return p


def add_table(headers, rows, col_widths, hdr_bg=NAVY, hdr_fg=WHITE,
              alt_bg=LTBLUE, font_size=8):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    # header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, hdr_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = hdr_fg
    # data rows
    for r_idx, row_data in enumerate(rows):
        bg = alt_bg if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
run = p.add_run('JAKEVOLUME')
run.bold = True
run.font.size = Pt(48)
run.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
run2 = p2.add_run('0DTE Options Alerting System')
run2.font.size = Pt(22)
run2.font.color.rgb = TEAL
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph()
run3 = p3.add_run('Architecture Reference')
run3.font.size = Pt(16)
run3.italic = True
run3.font.color.rgb = GOLD
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(30)

# summary box
tbl_cover = doc.add_table(rows=1, cols=3)
tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, val) in enumerate([
    ('Symbols', '9 (MAG7 + SPY + QQQ)'),
    ('Poll Interval', '60 seconds'),
    ('Signal Sources', 'Databento (8 AM) · Schwab (8:30 AM–3 PM)'),
]):
    c = tbl_cover.rows[0].cells[i]
    set_cell_bg(c, LTBLUE)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b = p.add_run(label + '\n')
    b.bold = True
    b.font.size = Pt(9)
    b.font.color.rgb = NAVY
    v = p.add_run(val)
    v.font.size = Pt(8)

p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ver.paragraph_format.space_before = Pt(40)
r = p_ver.add_run('Data: Databento Historical + Live  ·  Broker: Charles Schwab API  ·  Storage: PostgreSQL + Google Sheets')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — 8:00 AM MORNING SNAPSHOT
# ══════════════════════════════════════════════════════════════════════════════
hdr_para('8:00 AM CST  —  Morning Snapshot', size=18, color=NAVY)
sub_para('Runs ONCE per trading day.  Uses Databento Historical only — settled T+1 OI is the anchor for S/R levels.',
         italic=True, color=TEAL)

hdr_para('Trigger', size=11, color=TEAL)
add_table(
    ['Condition', 'Detail'],
    [
        ['Time window', '08:00–08:29 CST  (is_snapshot_window())'],
        ['Guard',       'snap_done != today — executes exactly once per calendar day'],
        ['Data source', 'Databento Historical (DatabentoClient)  —  Schwab NOT used here'],
    ],
    col_widths=[1.5, 8.0],
)

hdr_para('Data Collection  (Databento Historical)', size=11, color=TEAL)
add_table(
    ['Step', 'API / Method', 'Dataset', 'What it returns'],
    [
        ['1', 'dbc.get_prev_close(symbol)', 'XNAS.ITCH  ohlcv-1d', 'Previous session close price — ATM anchor for OI levels'],
        ['2', 'dbc.get_quote(symbol)',       'XNAS.ITCH  Live buffer', 'Pre-market price (falls back to prev_close if no bar yet)'],
        ['3', 'dbc.get_option_chain(symbol)','OPRA.PILLAR  mbp-10', 'Full option chain with settled OI from prior session'],
    ],
    col_widths=[0.35, 2.6, 2.2, 4.4],
)

hdr_para('Analysis', size=11, color=TEAL)
add_table(
    ['Function', 'Input', 'Output'],
    [
        ['compute_oi_levels(chain, prev_close)',
         'Option chain + ATM anchor',
         'R1/R2 = top call OI strikes closest to spot;  S1/S2 = top put OI strikes closest to spot'],
        ['get_top_oi_snapshot(chain, prev_close)',
         'Option chain + ATM anchor',
         'Top-2 call & put strikes sorted by raw OI (for OI_Snapshot sheet reference)'],
        ['compute_sentiment(chain, pm_price, prev_close)',
         'Chain + pre-market price',
         'Drift score + P/C OI ratio → Bias (BULLISH / NEUTRAL / BEARISH) + total_score'],
    ],
    col_widths=[2.8, 2.8, 3.9],
)

hdr_para('Outputs', size=11, color=TEAL)
add_table(
    ['Destination', 'Sheet / Table', 'Columns written'],
    [
        ['PostgreSQL', 'option_chains', 'symbol, snap_date, snap_time, expiry_date, contracts JSON, underlying_price'],
        ['PostgreSQL', 'oi_levels',     'symbol, date, computed_at, level_type, rank, strike, open_interest'],
        ['Google Sheets', 'Daily_Levels',     'Date, Symbol, Computed_At_CST, Underlying_Price, S1/S2 Strike+OI, R1/R2 Strike+OI'],
        ['Google Sheets', 'OI_Snapshot',      'Date, Time_CST, Symbol, Expiry, Call_1/2 Strike+OI, Put_1/2 Strike+OI, Underlying_Price'],
        ['Google Sheets', 'Morning_Sentiment','Date, Computed_At_CST, Symbol, Prev_Close, PM_Price, PM_Change_Pct, Call_OI, Put_OI, PC_Ratio, Drift_Score, PC_Score, Total_Score, Bias'],
        ['Console',    '─',              'MAG7 morning briefing table (symbol, prev_close, pm_price, change%, P/C, bias)'],
    ],
    col_widths=[1.5, 2.0, 6.0],
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — 8:30 AM INTRADAY LOOP
# ══════════════════════════════════════════════════════════════════════════════
hdr_para('8:30 AM – 3:00 PM CST  —  Intraday Signal Loop', size=18, color=NAVY)
sub_para('Runs every 60 seconds.  Uses Schwab for equity bars, expiry discovery, and option quotes.  Databento Live handles positioning monitor.',
         italic=True, color=TEAL)

hdr_para('Trigger', size=11, color=TEAL)
add_table(
    ['Condition', 'Detail'],
    [
        ['Time window',  '08:30–15:00 CST  (is_market_open())'],
        ['Frequency',    'Every 60 s — main loop sleeps for remainder of POLL_INTERVAL_SECONDS'],
        ['Primary source', 'SchwabClient — equity bars + option quotes'],
        ['Fallback',     'DatabentoClient — used if SchwabClient not initialised'],
    ],
    col_widths=[1.5, 8.0],
)

hdr_para('Data Collection  (Schwab — Primary)', size=11, color=TEAL)
add_table(
    ['Step', 'Method', 'What it returns'],
    [
        ['1', 'schwab.get_bars(symbol)',
              '1-min OHLCV bars — Price History API, today, regular session only.  Same dict format as DatabentoClient.'],
        ['2', 'schwab.get_nearest_expiry(symbol)',
              'Nearest date with both calls & puts.  Cached per symbol per trading day (1 API call/day).'],
        ['3', 'schwab.get_option_quotes_for_levels(symbol, expiry, levels)',
              'Real-time bid/ask/mark/volume for every S/R strike.  Keyed by (strike, option_type).'],
        ['4', 'db.get_today_levels(symbol, today)',
              'S/R levels computed at 8 AM from Postgres — used as strike list for quote fetch and signal proximity check.'],
    ],
    col_widths=[0.35, 3.1, 6.05],
)

hdr_para('Signal Detector  —  Conditions (all must be satisfied)', size=11, color=TEAL)
add_table(
    ['Condition', 'Threshold', 'Notes'],
    [
        ['Proximity to S/R level',   '≤ 0.5 % of strike price',       'Equity close price within 0.5% of an OI-derived S/R strike'],
        ['Equity volume spike',       '≥ 2 × 20-bar average volume',   'Current bar volume vs. rolling 20-bar mean'],
        ['Option delta volume',       '≥ 25 contracts on bid or ask',  'From Schwab real-time option quotes at that strike'],
        ['Adjusted cluster check',    'Consecutive bar confirmation',  'Volume cluster confirmed over ≥ 3 consecutive 1-min bars'],
        ['Signal cooldown',           '30 minutes per symbol',         'No duplicate signal within 30 min on the same symbol'],
    ],
    col_widths=[2.5, 2.5, 4.5],
)

hdr_para('Signal Semantics', size=11, color=TEAL)
add_table(
    ['Field', 'Source', 'Meaning'],
    [
        ['signal_type',      'SignalDetector',           'CALL_BREAKOUT or PUT_BREAKDOWN'],
        ['symbol',           'config.SYMBOLS',           'Underlying equity ticker (e.g. AAPL)'],
        ['option_type',      'Derived from signal_type', 'CALL or PUT'],
        ['level_price (strike)', 'OI level from Postgres', 'The S/R strike that triggered the signal'],
        ['expiry',           'schwab.get_nearest_expiry','0DTE or nearest available expiry date'],
        ['trigger_price',    'bars[-1]["close"]',        'Equity close price at signal bar — also Ticker_Price_At_Entry in Sheets'],
        ['price_to_enter',   'Schwab option ask price',  'Ask price of the option contract at signal time'],
        ['price_to_exit',    'Schwab option bid price',  'Bid price (immediate exit estimate)'],
        ['spike_volume',     'Current bar volume',       'Volume of the triggering equity bar'],
        ['signal_time',      'bar_time of trigger bar',  'CST timestamp of the 1-min bar that fired the signal'],
    ],
    col_widths=[1.8, 2.4, 5.3],
)

hdr_para('Outputs', size=11, color=TEAL)
add_table(
    ['Destination', 'Action', 'Detail'],
    [
        ['PostgreSQL',   'db.save_bars()',           'Stores all 1-min OHLCV bars each poll cycle'],
        ['PostgreSQL',   'db.save_signal()',         'Persists fired signal with all fields; returns signal_id'],
        ['PostgreSQL',   'db.mark_signal_logged()',  'Sets logged=True after Sheets write confirmed'],
        ['Google Sheets','sheets.log_signal()',      'Async enqueue → Signals sheet: Datetime_CST, Contract, Ticker_Price_At_Entry, Option_Price_To_Enter, Option_Price_To_Exit'],
        ['Desktop',      '_notify_signal() (plyer)', 'Optional toast notification: contract, enter/exit prices, spike volume'],
    ],
    col_widths=[1.5, 2.2, 5.8],
)

hdr_para('Positioning Monitor  (Databento Live — parallel)', size=11, color=TEAL)
add_table(
    ['Component', 'Detail'],
    [
        ['Source',   'dbc.get_expiry_pair(symbol) + dbc.get_atm_option_quotes_all_expiries(symbol, price)  via Databento Live'],
        ['Logic',    'monitor.update() detects SAME_DAY_MOVER and NEXT_EXPIRY_POSITIONING clusters'],
        ['Storage',  'Postgres only — no Sheets write, no signal fired; purely observational positioning data'],
        ['Isolation','Wrapped in try/except — failure does not interrupt signal detection for other symbols'],
    ],
    col_widths=[1.5, 8.0],
)

hdr_para('Key Config Parameters', size=11, color=TEAL)
add_table(
    ['Parameter', 'Default', 'Effect'],
    [
        ['POLL_INTERVAL_SECONDS', '60',  'Main loop cadence'],
        ['TOP_N_LEVELS',          '2',   'Number of S/R levels per side (R1/R2, S1/S2)'],
        ['BARS_TO_FETCH',         '30',  'Number of 1-min bars fetched per symbol per poll'],
        ['SYMBOLS',               '9',   'AAPL AMZN GOOGL META MSFT NVDA TSLA SPY QQQ'],
        ['SCHWAB_API_KEY',        '.env','Schwab developer app client ID'],
        ['SCHWAB_TOKEN_FILE',     'schwab_token.json', 'OAuth token cache — auto-refreshed on each run'],
        ['DATABENTO_API_KEY',     '.env','Databento Historical + Live auth key'],
        ['GOOGLE_SPREADSHEET_ID', '.env','Target Google Sheet for all log_* writes'],
    ],
    col_widths=[2.2, 1.8, 5.5],
)

out = r'C:\Users\malir\Projects\Python\Jakevolume\Jakevolume_Architecture.docx'
doc.save(out)
print(f'Saved: {out}')
