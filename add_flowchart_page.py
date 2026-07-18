"""
Append a two-column flowchart page to Jakevolume_Architecture.docx.
Left  : 8:00 AM Morning Snapshot
Right : 8:30 AM – 3:00 PM Intraday Signal Loop
"""
import os, tempfile
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Polygon
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Color palette ─────────────────────────────────────────────────────────────
C_NAVY    = '#1A375E'
C_TEAL    = '#007A87'
C_GOLD    = '#C88B00'
C_WHITE   = '#FFFFFF'
C_LTBLUE  = '#D6E4F0'
C_LTGREEN = '#D4EDDA'
C_LTGOLD  = '#FEF9E7'
C_LTORANGE= '#FAE5D3'
C_BG      = '#F0F4F8'
C_START   = '#1A375E'   # dark navy for start/end
C_TRIGGER = '#007A87'   # teal for triggers
C_DATA    = '#2E86AB'   # blue for data collection
C_COMPUTE = '#5B4E8A'   # purple for analysis/compute
C_STORE   = '#2D6A4F'   # green for storage
C_SIGNAL  = '#C8500A'   # orange for signal
C_DIAMOND = '#C88B00'   # gold for decisions


# ── Drawing helpers ───────────────────────────────────────────────────────────

def box(ax, cx, cy, w, h, lines, bg, fg='white', fs=7.5, bold=False,
        radius=0.12, lw=1.3):
    rect = FancyBboxPatch(
        (cx - w/2, cy - h/2), w, h,
        boxstyle=f'round,pad={radius}',
        facecolor=bg, edgecolor=_darken(bg), linewidth=lw,
        zorder=3
    )
    ax.add_patch(rect)
    text = '\n'.join(lines) if isinstance(lines, list) else lines
    ax.text(cx, cy, text, ha='center', va='center',
            fontsize=fs, color=fg,
            fontweight='bold' if bold else 'normal',
            multialignment='center', zorder=4,
            linespacing=1.35)


def diamond(ax, cx, cy, w, h, lines, bg=C_DIAMOND, fg='white', fs=7.5):
    pts = [(cx, cy+h/2), (cx+w/2, cy), (cx, cy-h/2), (cx-w/2, cy)]
    poly = Polygon(pts, closed=True, facecolor=bg,
                   edgecolor=_darken(bg), linewidth=1.3, zorder=3)
    ax.add_patch(poly)
    text = '\n'.join(lines) if isinstance(lines, list) else lines
    ax.text(cx, cy, text, ha='center', va='center',
            fontsize=fs, color=fg, fontweight='bold',
            multialignment='center', zorder=4)


def arr(ax, x1, y1, x2, y2, label='', side='right', color='#444444'):
    ax.annotate(
        '', xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle='->', color=color,
                        lw=1.4, mutation_scale=14),
        zorder=2
    )
    if label:
        mx = (x1+x2)/2 + (0.07 if side == 'right' else -0.07)
        my = (y1+y2)/2
        ax.text(mx, my, label, fontsize=6.5, color=color,
                ha='left' if side == 'right' else 'right', va='center')


def _darken(hex_color, factor=0.75):
    h = hex_color.lstrip('#')
    r, g, b = int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
    return f'#{int(r*factor):02X}{int(g*factor):02X}{int(b*factor):02X}'


def section_label(ax, cx, cy, text):
    ax.text(cx, cy, text, ha='center', va='center',
            fontsize=8, color='#555555', style='italic')


# ── Build figure ─────────────────────────────────────────────────────────────

fig = plt.figure(figsize=(20, 11))
fig.patch.set_facecolor(C_BG)

# Two axes side-by-side
ax1 = fig.add_axes([0.01, 0.04, 0.47, 0.88])   # left
ax2 = fig.add_axes([0.52, 0.04, 0.47, 0.88])   # right

for ax in (ax1, ax2):
    ax.set_xlim(0, 4)
    ax.set_ylim(0, 11)
    ax.axis('off')
    ax.set_facecolor(C_BG)

# Vertical divider
fig.add_artist(plt.Line2D([0.505, 0.505], [0.03, 0.97],
                           color='#AABBCC', linewidth=1.5, linestyle='--',
                           transform=fig.transFigure))

# ── COLUMN HEADERS ────────────────────────────────────────────────────────────
ax1.text(2, 10.65, '8:00 AM CST — Morning Snapshot',
         ha='center', va='center', fontsize=12,
         fontweight='bold', color=C_NAVY)
ax1.text(2, 10.35, 'Data source: Databento Historical only',
         ha='center', va='center', fontsize=8, color=C_TEAL, style='italic')

ax2.text(2, 10.65, '8:30 AM – 3:00 PM CST — Intraday Signal Loop',
         ha='center', va='center', fontsize=12,
         fontweight='bold', color=C_NAVY)
ax2.text(2, 10.35, 'Data source: Schwab (primary) · Databento Live (positioning monitor)',
         ha='center', va='center', fontsize=8, color=C_TEAL, style='italic')

# ════════════════════════════════════════════════════════
# LEFT COLUMN — 8:00 AM Morning Snapshot
# ════════════════════════════════════════════════════════
BW, BH = 2.6, 0.55   # default box width / height
CX = 2.0              # center x

Y = [9.8, 9.0, 8.15, 7.2, 6.35, 5.35, 4.4, 3.45, 2.55, 1.65, 0.85]

# START
box(ax1, CX, Y[0], 1.8, 0.48,
    'START — 08:00 CST', bg=C_START, bold=True, fs=8.5)

# Trigger
arr(ax1, CX, Y[0]-0.24, CX, Y[1]+0.27)
box(ax1, CX, Y[1], BW, BH,
    ['Trigger: is_snapshot_window()  +  snap_done ≠ today'],
    bg=C_TRIGGER, fs=7.5)

# Guard
arr(ax1, CX, Y[1]-0.27, CX, Y[2]+0.33)
diamond(ax1, CX, Y[2], 2.9, 0.62,
        ['snap_done == today?'], fs=7.5)
# "Yes → skip" branch to the right
arr(ax1, CX+1.45, Y[2], 3.8, Y[2], label='YES → skip', side='right', color='#888')
ax1.text(3.85, Y[2], 'wait for\nnext poll', ha='left', va='center',
         fontsize=6.5, color='#888888')

# Loop header
arr(ax1, CX, Y[2]-0.31, CX, Y[3]+0.27, label='NO')
box(ax1, CX, Y[3], BW, 0.48,
    ['For each symbol  (AAPL · AMZN · GOOGL · META · MSFT · NVDA · TSLA · SPY · QQQ)'],
    bg='#3D5A80', fs=6.8)

# Data collection
arr(ax1, CX, Y[3]-0.24, CX, Y[4]+0.33)
box(ax1, CX, Y[4], BW, 0.62,
    ['Databento Historical',
     'get_prev_close()  ·  get_quote()  ·  get_option_chain()'],
    bg=C_DATA, fs=7.5)
section_label(ax1, 3.75, Y[4], 'Data')

# Analysis
arr(ax1, CX, Y[4]-0.31, CX, Y[5]+0.33)
box(ax1, CX, Y[5], BW, 0.62,
    ['Analysis',
     'compute_oi_levels()  →  R1/R2 (calls) · S1/S2 (puts)\ncompute_sentiment()  →  Drift + P/C → BULL/BEAR/NEUT'],
    bg=C_COMPUTE, fs=7.0)
section_label(ax1, 3.75, Y[5], 'Analyse')

# Save Postgres
arr(ax1, CX, Y[5]-0.31, CX, Y[6]+0.27)
box(ax1, CX, Y[6], BW, BH,
    ['Save to PostgreSQL',
     'option_chains  ·  oi_levels'],
    bg=C_STORE, fs=7.5)
section_label(ax1, 3.75, Y[6], 'Persist')

# Log Sheets
arr(ax1, CX, Y[6]-0.27, CX, Y[7]+0.27)
box(ax1, CX, Y[7], BW, BH,
    ['Log to Google Sheets',
     'Daily_Levels  ·  OI_Snapshot  ·  Morning_Sentiment'],
    bg=C_STORE, fs=7.5)
section_label(ax1, 3.75, Y[7], 'Sheets')

# MAG7 briefing
arr(ax1, CX, Y[7]-0.27, CX, Y[8]+0.27)
box(ax1, CX, Y[8], BW, BH,
    ['Print MAG7 Console Briefing',
     'Symbol · Prev Close · PM Price · Bias · P/C Ratio'],
    bg='#5B4E8A', fs=7.5)

# snap_done = today
arr(ax1, CX, Y[8]-0.27, CX, Y[9]+0.22)
box(ax1, CX, Y[9], 2.0, 0.42,
    'snap_done = today', bg='#4A5568', fs=8)

# END
arr(ax1, CX, Y[9]-0.21, CX, Y[10]+0.22)
box(ax1, CX, Y[10], 1.8, 0.42,
    'END (wait for next poll)', bg=C_START, bold=True, fs=8)


# ════════════════════════════════════════════════════════
# RIGHT COLUMN — 8:30 AM Intraday Loop
# ════════════════════════════════════════════════════════
CX2 = 2.0
Y2 = [9.8, 9.0, 8.15, 7.25, 6.4, 5.55, 4.65, 3.8, 2.85, 1.85, 1.0, 0.3]

# START
box(ax2, CX2, Y2[0], 2.2, 0.48,
    '60-Second Poll Loop', bg=C_START, bold=True, fs=8.5)

# Market open?
arr(ax2, CX2, Y2[0]-0.24, CX2, Y2[1]+0.31)
diamond(ax2, CX2, Y2[1], 2.8, 0.58,
        ['is_market_open()?  08:30–15:00 CST'], fs=7.5)
arr(ax2, CX2+1.4, Y2[1], 3.8, Y2[1], label='NO', side='right', color='#888')
ax2.text(3.85, Y2[1], 'sleep 60s\n→ loop', ha='left', va='center',
         fontsize=6.5, color='#888888')

# Symbol loop
arr(ax2, CX2, Y2[1]-0.29, CX2, Y2[2]+0.27, label='YES')
box(ax2, CX2, Y2[2], BW, 0.48,
    ['For each symbol  (9 symbols)'],
    bg='#3D5A80', fs=7.5)

# Schwab bars + expiry
arr(ax2, CX2, Y2[2]-0.24, CX2, Y2[3]+0.33)
box(ax2, CX2, Y2[3], BW, 0.62,
    ['Schwab: get_bars()',
     '1-min OHLCV — regular session  ·  save to Postgres'],
    bg=C_DATA, fs=7.5)
section_label(ax2, 3.75, Y2[3], 'Bars')

# Expiry + quotes
arr(ax2, CX2, Y2[3]-0.31, CX2, Y2[4]+0.33)
box(ax2, CX2, Y2[4], BW, 0.62,
    ['Schwab: get_nearest_expiry()  (cached/day)',
     'Schwab: get_option_quotes_for_levels()  →  bid/ask/mark'],
    bg=C_DATA, fs=7.2)
section_label(ax2, 3.75, Y2[4], 'Quotes')

# Load S/R levels
arr(ax2, CX2, Y2[4]-0.31, CX2, Y2[5]+0.27)
box(ax2, CX2, Y2[5], BW, BH,
    ['db.get_today_levels()',
     'S/R levels computed at 8 AM — from Postgres'],
    bg='#4A7C59', fs=7.5)
section_label(ax2, 3.75, Y2[5], 'Levels')

# Signal detector
arr(ax2, CX2, Y2[5]-0.27, CX2, Y2[6]+0.33)
box(ax2, CX2, Y2[6], BW, 0.62,
    ['SignalDetector.check()',
     'Proximity ≤0.5%  ·  Vol ≥2×avg  ·  Opt Δvol ≥25\n3 consecutive bars  ·  30-min cooldown'],
    bg=C_COMPUTE, fs=7.0)
section_label(ax2, 3.75, Y2[6], 'Detect')

# Signal fired?
arr(ax2, CX2, Y2[6]-0.31, CX2, Y2[7]+0.31)
diamond(ax2, CX2, Y2[7], 2.4, 0.58,
        ['Signal\nFired?'], fs=8, bg=C_DIAMOND)

# NO branch → positioning monitor (right side offset)
arr(ax2, CX2+1.2, Y2[7], 3.5, Y2[7], label='NO', side='right', color='#888')
box(ax2, 3.7, Y2[8]-0.15, 0.55, 1.0,
    ['Position\nMonitor\n(Databento\nLive)'],
    bg='#607D8B', fs=6.0, radius=0.08)
arr(ax2, 3.7, Y2[8]-0.65, 3.7, Y2[10]+0.20, color='#888')

# YES branch — signal outputs
arr(ax2, CX2, Y2[7]-0.29, CX2, Y2[8]+0.27, label='YES')
box(ax2, CX2, Y2[8], BW, BH,
    ['db.save_signal()  →  db.mark_signal_logged()',
     'sheets.log_signal()  →  Signals sheet'],
    bg=C_SIGNAL, fs=7.5)
section_label(ax2, 3.75, Y2[8], 'Log')

# Google Sheets signal details
arr(ax2, CX2, Y2[8]-0.27, CX2, Y2[9]+0.33)
box(ax2, CX2, Y2[9], BW, 0.62,
    ['Signals Sheet Row',
     'Datetime_CST · Contract · Ticker_Price_At_Entry\nOption_Price_To_Enter · Option_Price_To_Exit'],
    bg='#2D6A4F', fs=7.0)
section_label(ax2, 3.75, Y2[9], 'Sheets')

# Desktop notify
arr(ax2, CX2, Y2[9]-0.31, CX2, Y2[10]+0.22)
box(ax2, CX2, Y2[10], BW, 0.42,
    ['Desktop Notification  (plyer — optional)'],
    bg='#6C3483', fs=7.5)

# Sleep / loop back
arr(ax2, CX2, Y2[10]-0.21, CX2, Y2[11]+0.15, color='#444')
box(ax2, CX2, Y2[11], 2.2, 0.28,
    'Sleep → next 60s poll', bg='#4A5568', fs=8)

# Loop-back arrow (left side of right column)
arr(ax2, CX2-1.1, Y2[11], 0.1, Y2[11], color='#AAAAAA')
ax2.annotate('', xy=(0.1, Y2[0]), xytext=(0.1, Y2[11]),
             arrowprops=dict(arrowstyle='->', color='#AAAAAA', lw=1.3))
ax2.text(0.03, (Y2[0]+Y2[11])/2, 'loop', ha='center', va='center',
         fontsize=6.5, color='#AAAAAA', rotation=90)

# ── Legend ────────────────────────────────────────────────────────────────────
legend_items = [
    (C_TRIGGER, 'Trigger / Condition'),
    (C_DATA,    'Data Fetch'),
    (C_COMPUTE, 'Analysis / Compute'),
    (C_STORE,   'Storage (Postgres)'),
    (C_SIGNAL,  'Signal Output'),
    (C_DIAMOND, 'Decision'),
]
lx, ly = 0.01, 0.015
for i, (color, label) in enumerate(legend_items):
    xpos = lx + i * 0.155
    rect = mpatches.Patch(facecolor=color, edgecolor=_darken(color), label=label)
    fig.text(xpos + 0.012, ly, '■', color=color, fontsize=11,
             transform=fig.transFigure, va='center')
    fig.text(xpos + 0.025, ly, label, fontsize=7.2,
             transform=fig.transFigure, va='center', color='#333333')

fig.text(0.5, 0.005, 'Jakevolume 0DTE Alerting System — System Architecture Flow',
         ha='center', fontsize=8, color='#666666', style='italic')

# ── Save PNG ──────────────────────────────────────────────────────────────────
tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
tmp.close()
fig.savefig(tmp.name, dpi=180, bbox_inches='tight', facecolor=C_BG)
plt.close(fig)
print(f'Flowchart saved to temp: {tmp.name}')

# ── Append to existing docx ───────────────────────────────────────────────────
docx_path = r'C:\Users\malir\Projects\Python\Jakevolume\Jakevolume_Architecture.docx'
doc = Document(docx_path)

# Page break to start new page
doc.add_page_break()

# Page heading
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('System Architecture — Flow Diagrams')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1A, 0x37, 0x5E)
p.paragraph_format.space_after = Pt(6)

# Insert flowchart image — fill the page width
doc.add_picture(tmp.name, width=Inches(9.8))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

os.unlink(tmp.name)

doc.save(docx_path)
print(f'Updated: {docx_path}')
